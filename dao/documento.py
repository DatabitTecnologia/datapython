import json

from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import log
import pandas as pd
import connect
import datetime
import base64
from docx import Document as docx
import os
from pathlib import Path
from docx2pdf import convert

wdFormatPDF = 17
import win32com.client

class Documento:

    def move_table_after(self, table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def set_col_widths(self, table, widthcolumns):
        for row in table.rows:
            c = 0
            fator_conversao = 0.10
            for col in table.columns:
                row.cells[c].width = Inches(widthcolumns[c] * fator_conversao)
                c += 1

    def set_cell_border(self, cell, **kwargs):
        """
            Set cell`s border
            Usage:

            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                bottom={"sz": 12, "color": "#00FF00", "val": "single"},
                start={"sz": 24, "val": "dashed", "shadow": "true"},
                end={"sz": 12, "val": "dashed"},
            )
            """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def createDocument(self, params):
        try:
            df_params = pd.DataFrame.from_dict(params, orient="index")

            stoken = str(df_params.loc["token", 0])

            try:
                smodel = str(df_params.loc["model", 0])
            except Exception:
                smodel = ""

            try:
                smov = str(df_params.loc["mov", 0])
            except Exception:
                smov = ""

            try:
                suser = str(df_params.loc["user", 0])
            except Exception:
                suser = ""

            dados = connect.dataconnect(stoken, 0)

            # Dados Proposta
            cursor_proposta = dados.cursor()
            cursor_proposta.execute("SELECT * FROM TB01140 WHERE TB01140_CODIGO = '{}' ".format(smodel))
            valores_proposta = cursor_proposta.fetchall()
            descricao_proposta = cursor_proposta.description
            colunas_proposta = [tp[0] for tp in descricao_proposta]
            campos_proposta = pd.DataFrame.from_records(valores_proposta, columns=colunas_proposta)
            ioperacao = campos_proposta["TB01140_OPERACAO"][0]
            sarquivofim = "c:\\temp\\" + campos_proposta["TB01140_ARQUIVOFIM"][0]
            sarquivofim2 = campos_proposta["TB01140_ARQUIVOFIM"][0]
            sarquivo = suser + '_' + smov + '_' + campos_proposta["TB01140_ARQUIVO"][0]
            stitulo = campos_proposta["TB01140_TITULO"][0]
            sobs = campos_proposta["TB01140_OBS"][0]

            # Base64 do arquivo
            cursor_file = dados.cursor()
            ssqlfile = "SELECT REPLACE(REPLACE((SELECT CAST(TB01140_BYTES as VARBINARY(MAX)) as ARQ " \
                       "FROM TB01140 WHERE TB01140_CODIGO = '{}' FOR XML  PATH(''), BINARY BASE64)," \
                       "'<ARQ>',''),'</ARQ>','') AS arq;".format(smodel)
            cursor_file.execute(ssqlfile)
            valores_file = cursor_file.fetchall()
            descricao_file = cursor_file.description
            colunas_file = [tp[0] for tp in descricao_file]
            campos_file = pd.DataFrame.from_records(valores_file, columns=colunas_file)
            sbase64 = campos_file["arq"][0]

            # Criando arquivo tempor?rio
            bytesarq = base64.b64decode(sbase64, validate=True)
            newfile = "c:\\temp\\" + sarquivo
            f = open(newfile, 'wb')
            f.write(bytesarq)
            f.close()

            # Abrindo documento tempor?rio
            documento = docx("c:\\temp\\" + sarquivo)

            # Dados Movimento
            if int(ioperacao) == 0:
                stabela = "TB02255"
            elif int(ioperacao) == 3:
                stabela = "TB02264"
            else:
                stabela = "TB02303"

            cursor_movimento = dados.cursor()
            cursor_movimento.execute("SELECT * FROM {} WHERE {}_CODIGO = '{}' ".format(stabela, stabela, smov))
            valores_movimento = cursor_movimento.fetchall()
            descricao_movimento = cursor_movimento.description
            colunas_movimento = [tp[0] for tp in descricao_movimento]
            campos_movimento = pd.DataFrame.from_records(valores_movimento, columns=colunas_movimento)

            # Dados PlaceHolder
            cursor_placeholder = dados.cursor()
            cursor_placeholder.execute("SELECT * FROM TB01141 WHERE TB01141_MODELO = '{}' ".format(smodel))
            valores_placeholder = cursor_placeholder.fetchall()
            descricao_placeholder = cursor_placeholder.description
            colunas_placeholder = [tp[0] for tp in descricao_placeholder]
            campos_placeholder = pd.DataFrame.from_records(valores_placeholder, columns=colunas_placeholder)

            cursor_item = dados.cursor()

            for i, placeholder in campos_placeholder.iterrows():
                splaceholder = placeholder["TB01141_PLACEHOLDER"]
                scampo = placeholder["TB01141_CAMPO"]
                svalor = placeholder["TB01141_VALOR"]
                itipo = int(placeholder["TB01141_TIPO"])
                itipotabela = int(placeholder["TB01141_TIPOTABELA"])
                ssql = placeholder["TB01141_SQL"]
                itamanho = int(placeholder["TB01141_TAMANHO"])
                stabela = placeholder["TB01141_TABELA"]
                scampochave = placeholder["TB01141_CAMPOCHAVE"]
                scamporef = placeholder["TB01141_CAMPOREF"]
                svalorfim = ""

                if ssql is None:
                    ssql = ""
                if itipo == 0:  # Definir campo
                    svalorfim = str(campos_movimento[scampo][0])[0:itamanho]
                elif itipo == 1:  # Valor Fixo
                    svalorfim = svalor
                elif itipo == 2:  # Tabela Referenciada
                    if itipotabela == 0:  # Individual ou listagem simples"
                        cursor_item.execute(
                            "SELECT {} FROM {} WHERE {} = '{}' {}".format(scamporef, stabela, scampochave,
                                                                          campos_movimento[scampo][0], ssql))
                        valores_item = cursor_item.fetchall()
                        descricao_item = cursor_item.description
                        colunas_item = [tp[0] for tp in descricao_item]
                        campos_item = pd.DataFrame.from_records(valores_item, columns=colunas_item)
                        if len(valores_item) > 0:
                            svaloritem = campos_item[scamporef][0]
                            if itipotabela == 0:
                                svalorfim = svaloritem
                    else:
                        if scampo != "" and scampo is not None:
                            svalorcampo = campos_movimento[scampo][0]
                            if ssql != "" and ssql is not None:
                                ssql = ssql.replace("@#{}#@".format(scampo), svalorcampo)
                        cursor_item.execute(ssql)
                        valores_item = cursor_item.fetchall()
                        descricao_item = cursor_item.description
                        colunas_item = [tp[0] for tp in descricao_item]
                        campos_item = pd.DataFrame.from_records(valores_item, columns=colunas_item)

                elif itipo == 3:  # Data Atual
                    data = datetime.date.today()
                    svalorfim = "{}/{}/{}".format(data.day, data.month, data.year)
                elif itipo == 4:  # Data em Extenso
                    data = datetime.date.today()
                    if data.month == 1:
                        smes = "Janeiro"
                    elif data.month == 2:
                        smes = "Fevereiro"
                    elif data.month == 3:
                        smes = "Marco"
                    elif data.month == 4:
                        smes = "Abtil"
                    elif data.month == 5:
                        smes = "Maio"
                    elif data.month == 6:
                        smes = "Junho"
                    elif data.month == 7:
                        smes = "Julho"
                    elif data.month == 8:
                        smes = "Agosto"
                    elif data.month == 9:
                        smes = "Setembro"
                    elif data.month == 10:
                        smes = "Outubro"
                    elif data.month == 11:
                        smes = "Novembro"
                    elif data.month == 12:
                        smes = "Dezembro"
                    svalorfim = "{} de {} de {}".format(data.day, smes, data.year)

                if svalorfim is None:
                    svalorfim = ""

                # Alterando informa??es
                sarquivofim = sarquivofim.replace(splaceholder, str(svalorfim))
                sarquivofim2 = sarquivofim2.replace(splaceholder, str(svalorfim))
                stitulo = stitulo.replace(splaceholder, str(svalorfim))
                sobs = sobs.replace(splaceholder, str(svalorfim))
                sarquivodoc = sarquivofim + '.docx'
                sarquivopdf = sarquivofim + '.pdf'

                if itipo != 2:
                    for paragrafo in documento.paragraphs:
                        paragrafo.text = paragrafo.text.replace(splaceholder, str(svalorfim))
                else:
                    if itipotabela != 2:
                        for paragrafo in documento.paragraphs:
                            paragrafo.text = paragrafo.text.replace(splaceholder, str(svalorfim))
                    else:
                        if len(campos_item) > 0:
                            rows = []
                            for x, item in campos_item.iterrows():
                                row = {}
                                for coluna in colunas_item:
                                    row[coluna] = item[coluna]
                                rows.append(row)

                            table = documento.add_table(
                                rows=len(campos_item) + 1, cols=len(colunas_item)
                            )

                            c = 0
                            for coluna in colunas_item:
                                table.cell(0, c).text = coluna
                                c += 1
                            widthcolumns = []
                            for i, row in enumerate(rows):
                                x = 0
                                for chave in row.keys():
                                    table.cell(i + 1, x).text = str(row[chave])
                                    widthcolumns.append(len(row[chave]))
                                    x += 1

                            table.autofit = False
                            table.allow_autofit = False
                            self.set_col_widths(table, widthcolumns)

                            for cell in table._cells:
                                self.set_cell_border(
                                    cell,
                                    top={"sz": 3, "val": "single"},
                                    bottom={"sz": 3, "val": "single"},
                                    start={"sz": 3, "val": "single"},
                                    end={"sz": 3, "val": "single"},
                                )

                        for paragrafo in documento.paragraphs:
                            if splaceholder in paragrafo.text:
                                paragrafo.text = paragrafo.text.replace(splaceholder, "")
                                if len(campos_item) > 0:
                                    self.move_table_after(table, paragrafo)

            # Salvando o doc tempor?rio
            documento.save(sarquivodoc)

            # Gerando o PDF
            #convert(sarquivodoc, sarquivopdf)

            # Buscando tamanho do arquivo
            itamanho = os.path.getsize(sarquivodoc)

            # Fechando Conex?es
            cursor_file.close()
            cursor_proposta.close()
            cursor_movimento.close()
            cursor_item.close()
            cursor_placeholder.close()
            dados.close()

            # Gerando Base64 do PDF
            file_text = open(sarquivodoc, 'rb')
            file_read = file_text.read()
            file_encode = str(base64.encodebytes(file_read))
            file_encode = file_encode.replace("b'", "")
            file_encode = file_encode.replace("\\n", "")
            file_encode = file_encode.replace("'", "")

            # Gerando Retorno
            retorno = {"arquivo": sarquivofim2 + '.docx', "tamanho": itamanho, "titulo": stitulo, "mensagem": sobs,
                       "base64": str(file_encode)}

            return json.dumps(retorno, indent=2)
        except Exception as e:
            return log.geralog(-1, 'Erro', "Problemas na geracao do arquivo: {}".format(e))

    def converttoPDF(self, params):
        try:
            df_params = pd.DataFrame.from_dict(params, orient="index")
            try:
                sarquivodoc = str(df_params.loc["arquivodoc", 0])
            except Exception:
                sarquivodoc = ""

            try:
                arq64 = str(df_params.loc["arq64", 0])
            except Exception:
                arq64 = ""


            bytesarq = base64.b64decode(arq64, validate=True)
            newfile = sarquivodoc
            f = open(newfile, 'wb')
            f.write(bytesarq)
            f.close()

            sarquivofim = sarquivodoc.replace(".docx", ".pdf")
            sarquivopdf = sarquivofim

            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            wdFormatPDF = 17
            docx_filepath = Path(newfile).resolve()
            pdf_filepath = Path(sarquivopdf).resolve()
            doc = word.Documents.Open(str(docx_filepath))
            doc.SaveAs(str(pdf_filepath), FileFormat=wdFormatPDF)
            doc.Close()


            # Buscando tamanho do arquivo
            itamanho = os.path.getsize(sarquivopdf)

            # Gerando Base64 do PDF
            file_text = open(sarquivopdf, 'rb')
            file_read = file_text.read()
            file_encode = str(base64.encodebytes(file_read))
            file_encode = file_encode.replace("b'", "")
            file_encode = file_encode.replace("\\n", "")
            file_encode = file_encode.replace("'", "")

            # Gerando Retorno
            retorno = {"arquivo": sarquivofim , "tamanho": itamanho, "base64": str(file_encode)}

            return json.dumps(retorno, indent=2)
        except Exception as e:
            return log.geralog(-1, 'Erro', "Problemas na geracao do arquivo: {}".format(e))


    def convertDocument(self, params):
        try:
            df_params = pd.DataFrame.from_dict(params, orient="index")

            stoken = str(df_params.loc["token", 0])

            try:
                smodel = str(df_params.loc["model", 0])
            except Exception:
                smodel = ""

            try:
                smov = str(df_params.loc["mov", 0])
            except Exception:
                smov = ""

            try:
                suser = str(df_params.loc["user", 0])
            except Exception:
                suser = ""

            dados = connect.dataconnect(stoken, 0)

            # Dados Proposta
            cursor_proposta = dados.cursor()
            cursor_proposta.execute("SELECT * FROM TB01140 WHERE TB01140_CODIGO = '{}' ".format(smodel))
            valores_proposta = cursor_proposta.fetchall()
            descricao_proposta = cursor_proposta.description
            colunas_proposta = [tp[0] for tp in descricao_proposta]
            campos_proposta = pd.DataFrame.from_records(valores_proposta, columns=colunas_proposta)
            ioperacao = campos_proposta["TB01140_OPERACAO"][0]
            sarquivofim = "c:\\temp\\" + campos_proposta["TB01140_ARQUIVOFIM"][0]
            sarquivofim2 = campos_proposta["TB01140_ARQUIVOFIM"][0]
            sarquivo = suser + '_' + smov + '_' + campos_proposta["TB01140_ARQUIVO"][0]
            stitulo = campos_proposta["TB01140_TITULO"][0]
            sobs = campos_proposta["TB01140_OBS"][0]

            # Base64 do arquivo
            cursor_file = dados.cursor()
            ssqlfile = "SELECT REPLACE(REPLACE((SELECT CAST(TB01140_BYTES as VARBINARY(MAX)) as ARQ " \
                       "FROM TB01140 WHERE TB01140_CODIGO = '{}' FOR XML  PATH(''), BINARY BASE64)," \
                       "'<ARQ>',''),'</ARQ>','') AS arq;".format(smodel)
            cursor_file.execute(ssqlfile)
            valores_file = cursor_file.fetchall()
            descricao_file = cursor_file.description
            colunas_file = [tp[0] for tp in descricao_file]
            campos_file = pd.DataFrame.from_records(valores_file, columns=colunas_file)
            sbase64 = campos_file["arq"][0]

            # Criando arquivo tempor?rio
            bytesarq = base64.b64decode(sbase64, validate=True)
            newfile = "c:\\temp\\" + sarquivo
            f = open(newfile, 'wb')
            f.write(bytesarq)
            f.close()

            # Abrindo documento tempor?rio
            #documento = docx("c:\\temp\\" + sarquivo)

            # Dados Movimento
            if int(ioperacao) < 3:
                stabela = "TB02255"
            else:
                stabela = "TB02264"
            cursor_movimento = dados.cursor()
            cursor_movimento.execute("SELECT * FROM {} WHERE {}_CODIGO = '{}' ".format(stabela, stabela, smov))
            valores_movimento = cursor_movimento.fetchall()
            descricao_movimento = cursor_movimento.description
            colunas_movimento = [tp[0] for tp in descricao_movimento]
            campos_movimento = pd.DataFrame.from_records(valores_movimento, columns=colunas_movimento)

            # Dados PlaceHolder
            cursor_placeholder = dados.cursor()
            cursor_placeholder.execute("SELECT * FROM TB01141 WHERE TB01141_MODELO = '{}' ".format(smodel))
            valores_placeholder = cursor_placeholder.fetchall()
            descricao_placeholder = cursor_placeholder.description
            colunas_placeholder = [tp[0] for tp in descricao_placeholder]
            campos_placeholder = pd.DataFrame.from_records(valores_placeholder, columns=colunas_placeholder)

            cursor_item = dados.cursor()
            doc = docx(newfile)
            for i, placeholder in campos_placeholder.iterrows():
                splaceholder = placeholder["TB01141_PLACEHOLDER"]
                print(splaceholder)
                scampo = placeholder["TB01141_CAMPO"]
                svalor = placeholder["TB01141_VALOR"]
                itipo = int(placeholder["TB01141_TIPO"])
                itipotabela = int(placeholder["TB01141_TIPOTABELA"])
                ssql = placeholder["TB01141_SQL"]
                itamanho = int(placeholder["TB01141_TAMANHO"])
                stabela = placeholder["TB01141_TABELA"]
                scampochave = placeholder["TB01141_CAMPOCHAVE"]
                scamporef = placeholder["TB01141_CAMPOREF"]
                svalorfim = ""

                if ssql is None:
                    ssql = ""
                if itipo == 0:  # Definir campo
                    svalorfim = str(campos_movimento[scampo][0])[0:itamanho]
                elif itipo == 1:  # Valor Fixo
                    svalorfim = svalor
                elif itipo == 2:  # Tabela Referenciada
                    if itipotabela == 0:  # Individual ou listagem simples"
                        cursor_item.execute(
                            "SELECT {} FROM {} WHERE {} = '{}' {}".format(scamporef, stabela, scampochave,
                                                                          campos_movimento[scampo][0], ssql))
                        valores_item = cursor_item.fetchall()
                        descricao_item = cursor_item.description
                        colunas_item = [tp[0] for tp in descricao_item]
                        campos_item = pd.DataFrame.from_records(valores_item, columns=colunas_item)
                        if len(valores_item) > 0:
                            svaloritem = campos_item[scamporef][0]
                            print(svaloritem)
                            if itipotabela == 0:
                                svalorfim = svaloritem

                    else:
                        if scampo != "" and scampo is not None:
                            svalorcampo = campos_movimento[scampo][0]
                            if ssql != "" and ssql is not None:
                                try:
                                    ssql = ssql.replace("@#{}#@".format(scampo), svalorcampo)
                                except:
                                    ssql = ssql.replace("@#{}#@".format(scampo), str(svalorcampo))

                        cursor_item.execute(ssql)
                        valores_item = cursor_item.fetchall()
                        descricao_item = cursor_item.description
                        colunas_item = [tp[0] for tp in descricao_item]
                        campos_item = pd.DataFrame.from_records(valores_item, columns=colunas_item)
                        if itipotabela != 2:
                            if valores_item:  # Verifica se há resultados
                                primeiro_campo = valores_item[0][0]
                                svalorfim = primeiro_campo
                            else:
                                svalorfim = ""


                elif itipo == 3:  # Data Atual
                    data = datetime.date.today()
                    svalorfim = "{}/{}/{}".format(data.day, data.month, data.year)
                elif itipo == 4:  # Data em Extenso
                    data = datetime.date.today()
                    if data.month == 1:
                        smes = "Janeiro"
                    elif data.month == 2:
                        smes = "Fevereiro"
                    elif data.month == 3:
                        smes = "Marco"
                    elif data.month == 4:
                        smes = "Abtil"
                    elif data.month == 5:
                        smes = "Maio"
                    elif data.month == 6:
                        smes = "Junho"
                    elif data.month == 7:
                        smes = "Julho"
                    elif data.month == 8:
                        smes = "Agosto"
                    elif data.month == 9:
                        smes = "Setembro"
                    elif data.month == 10:
                        smes = "Outubro"
                    elif data.month == 11:
                        smes = "Novembro"
                    elif data.month == 12:
                        smes = "Dezembro"
                    svalorfim = "{} de {} de {}".format(data.day, smes, data.year)

                if svalorfim is None:
                    svalorfim = ""
                print(svalorfim)
                # Alterando informa??es
                sarquivofim = sarquivofim.replace(splaceholder, str(svalorfim))
                sarquivofim2 = sarquivofim2.replace(splaceholder, str(svalorfim))
                stitulo = stitulo.replace(splaceholder, str(svalorfim))
                sobs = sobs.replace(splaceholder, str(svalorfim))
                sarquivodoc = sarquivofim2 + '.docx'
                sarquivopdf = sarquivofim2 + '.pdf'

                if itipo != 2:
                    for paragrafo in doc.paragraphs:
                        for run in paragrafo.runs:
                            if splaceholder in run.text:
                                run.text = run.text.replace(splaceholder, str(svalorfim))
                                font = run.font
                                run.font.name = font.name  # Manter o nome da fonte
                                run.font.size = font.size  # Manter o tamanho da fonte
                                run.font.bold = font.bold  # Manter o negrito, se aplicável
                                run.font.italic = font.italic  # Manter o itálico, se aplicável
                                run.font.color.rgb = font.color.rgb  # Manter a cor do texto
                else:
                    if itipotabela != 2:
                        for paragrafo in doc.paragraphs:
                            for run in paragrafo.runs:
                                if splaceholder in run.text:
                                   run.text = run.text.replace(splaceholder, str(svalorfim))
                                   font = run.font
                                   run.font.name = font.name  # Manter o nome da fonte
                                   run.font.size = font.size  # Manter o tamanho da fonte
                                   run.font.bold = font.bold  # Manter o negrito, se aplicável
                                   run.font.italic = font.italic  # Manter o itálico, se aplicável
                                   run.font.color.rgb = font.color.rgb  # Manter a cor do texto
                    else:
                        if len(campos_item) > 0:
                            rows = []
                            for x, item in campos_item.iterrows():
                                row = {}
                                for coluna in colunas_item:
                                    row[coluna] = item[coluna]
                                rows.append(row)

                            table = doc.add_table(
                                rows=len(campos_item) + 1, cols=len(colunas_item)
                            )

                            c = 0
                            for coluna in colunas_item:
                                table.cell(0, c).text = coluna
                                c += 1
                            widthcolumns = []
                            for i, row in enumerate(rows):
                                x = 0
                                for chave in row.keys():
                                    table.cell(i + 1, x).text = str(row[chave])
                                    widthcolumns.append(len(row[chave]))
                                    x += 1

                            table.autofit = False
                            table.allow_autofit = False
                            self.set_col_widths(table, widthcolumns)

                            for cell in table._cells:
                                self.set_cell_border(
                                    cell,
                                    top={"sz": 3, "val": "single"},
                                    bottom={"sz": 3, "val": "single"},
                                    start={"sz": 3, "val": "single"},
                                    end={"sz": 3, "val": "single"},
                                )

                        for paragrafo in doc.paragraphs:
                            if splaceholder in paragrafo.text:
                                for run in paragrafo.runs:
                                    if splaceholder in run.text:
                                        run.text = run.text.replace(splaceholder, "")
                                        font = run.font
                                        run.font.name = font.name  # Manter o nome da fonte
                                        run.font.size = font.size  # Manter o tamanho da fonte
                                        run.font.bold = font.bold  # Manter o negrito, se aplicável
                                        run.font.italic = font.italic  # Manter o itálico, se aplicável
                                        run.font.color.rgb = font.color.rgb  # Manter a cor do texto
                                        if len(campos_item) > 0:
                                            self.move_table_after(table, paragrafo)


            doc.save(sarquivodoc)


            #word = win32com.client.Dispatch('Word.Application')
            #docsave = word.Documents.Open(newfile)
            #docsave.SaveAs(sarquivopdf, FileFormat=17)
            #docsave.Close()
            #print(1)
            #convert(newfile, sarquivopdf)
            #print(2)

            itamanho = os.path.getsize(sarquivodoc)

            # Gerando Base64 do PDF
            file_text = open(sarquivodoc, 'rb')
            file_read = file_text.read()
            file_encode = str(base64.encodebytes(file_read))
            file_encode = file_encode.replace("b'", "")
            file_encode = file_encode.replace("\\n", "")
            file_encode = file_encode.replace("'", "")

            retorno = {"arquivo": sarquivodoc, "tamanho": itamanho, "base64": str(file_encode)}

            return json.dumps(retorno, indent=2)
        except Exception as e:
            print(e)
            return log.geralog(-1, 'Erro', "Problemas na geracao do arquivo: {}".format(e))
